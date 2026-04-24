#!/usr/bin/env python3
"""Generate docs/PRESENTATION.docx — team prep doc for Athernex 2026.

Mirrors the structure of README.md but formatted as a Word document the team
can print, highlight, and carry into the venue. Uses only python-docx so it
runs cleanly in CI and on a fresh VM.

Usage:
    python scripts/generate_presentation_docx.py

Writes to: docs/PRESENTATION.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "docs" / "PRESENTATION.docx"

BRAND = RGBColor(0x6E, 0x56, 0xCF)  # purple
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x52, 0x5F, 0x7A)
GREEN = RGBColor(0x15, 0x80, 0x3D)
RED = RGBColor(0xB4, 0x2A, 0x2A)


def _shade(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND if level <= 1 else INK


def add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 11,
    color: RGBColor | None = None,
    align: int | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "6E56CF")
    for r_i, row in enumerate(rows, start=1):
        cells = table.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = val
            for para in cells[c_i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
            cells[c_i].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def pagebreak(doc: Document) -> None:
    doc.add_page_break()


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── Cover ────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("VendorGuard AI")
    tr.bold = True
    tr.font.size = Pt(44)
    tr.font.color.rgb = BRAND

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "The Vendor Access Control Plane for DPDP-compliant India"
    )
    sr.italic = True
    sr.font.size = Pt(14)
    sr.font.color.rgb = INK

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Presentation prep · Athernex 2026 (DSCE × BMSCE)\n"
        "Team Rashi Innovators · v3.2.1"
    )
    mr.font.size = Pt(11)
    mr.font.color.rgb = MUTED

    add_para(doc, "", size=2)
    add_para(
        doc,
        "Every enterprise in India runs on vendors, and under the DPDP Act "
        "any one vendor breach is YOUR ₹250-crore problem. VendorGuard AI "
        "catches it in 43 milliseconds — with evidence, not vibes.",
        italic=True,
        size=13,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    pagebreak(doc)

    # ── Table of contents ────────────────────────────────────────────────────
    add_heading(doc, "Contents", level=1)
    toc = [
        "1. The problem — why DPDP makes this urgent",
        "2. Our solution in 30 seconds",
        "3. Why we picked this idea",
        "4. Real-world impact scenarios",
        "5. 5-layer architecture",
        "6. Tech stack — languages, frameworks, services",
        "7. What we cover that others miss",
        "8. Competitive landscape",
        "9. Business model — how we make money",
        "10. Killer pitch lines (memorize these)",
        "11. 90-second pitch delivery script",
        "12. Judge Q&A — 15 hard questions with answers",
        "13. Glossary — every abbreviation + full form",
        "14. Team knowledge base & stage-day checklist",
        "15. Changelog (v2 → v3.2.1)",
        "16. Credits & license",
    ]
    for t in toc:
        doc.add_paragraph(t)

    pagebreak(doc)

    # ── 1. Problem ────────────────────────────────────────────────────────────
    add_heading(doc, "1. The problem — why DPDP makes this urgent", level=1)
    add_para(
        doc,
        "India just made data breaches a ₹250 crore per-instance problem.",
        bold=True,
        size=13,
    )
    add_bullets(
        doc,
        [
            "The Digital Personal Data Protection Act 2023 (DPDP Act) is in force. "
            "The DPDP Rules 2025 (R.1–R.22) are landing now.",
            "Every company handling personal data of an Indian resident is a "
            "Data Fiduciary and is accountable not just for their own practices "
            "but for every vendor they hand PII to (§8(5)).",
            "Penalties: up to ₹250 crore per type of contravention (Schedule). "
            "A single vendor breach can trigger 3–5 contraventions at once — "
            "combined exposure ₹500 Cr – ₹1,250 Cr.",
            "Breach-notification SLA to the Data Protection Board: 72 hours (§8(6)).",
            "CERT-In 2022 directions require 6-hour cyber-incident reporting.",
        ],
    )
    add_para(doc, "", size=4)
    add_para(
        doc,
        "The existing compliance stack is built for GDPR. OneTrust, Securiti "
        "and BigID sell $50K–$80K/year GDPR-era SaaS with bolt-on DPDP "
        "modules. For a 200-vendor SME in Bangalore, that's a non-starter. "
        "Tsaaro, Cerebrus, Lexplosion ship PDFs, not a runtime product. "
        "No one is protecting the middle of the Indian market. That is our "
        "opening.",
        size=11,
    )

    # ── 2. Solution ──────────────────────────────────────────────────────────
    add_heading(doc, "2. Our solution in 30 seconds", level=1)
    add_para(
        doc,
        "VendorGuard AI is a one-screen, evidence-grounded, DPDP-native vendor "
        "access control plane.",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "Rules where rules belong — contract analysis, clause mapping.",
            "ML where ML belongs — behavioural anomaly detection at the gateway.",
            "LLM only for language polish — Executive summary, optional, always "
            "grounded in RAG.",
            "Nothing hidden. Every finding is traceable to a keyword, an offset, "
            "a gazette page, a rupee penalty, and a framework control.",
        ],
    )

    # ── 3. Why this idea ─────────────────────────────────────────────────────
    add_heading(doc, "3. Why we picked this idea", level=1)
    add_table(
        doc,
        ["Reason", "Why it matters"],
        [
            [
                "Live regulation, zero mature tooling",
                "DPDP Act is in force this year. First-mover window is open for "
                "~18 months before OneTrust ships a credible India module.",
            ],
            [
                "Clear, quantifiable pain",
                "₹250 Cr/contravention is board-agenda material. CFOs and "
                "CISOs already have budget; no market education needed.",
            ],
            [
                "Underserved India vendor long-tail",
                "100K+ SMEs with ≥50 vendors each. Huge greenfield below the "
                "$50K OneTrust floor.",
            ],
            [
                "Evidence-first fits Indian regulatory culture",
                "RBI, SEBI, DPB want paper trails. Our auditable JSON + "
                "gazette citations + CERT-In PDFs are what regulators ask for.",
            ],
            [
                "Demoable in 90 seconds",
                "Trust ring animates, gateway fires sub-100ms, Audit ZIP "
                "downloads live. Hackathon-optimised without sacrificing depth.",
            ],
            [
                "We could actually build it",
                "FastAPI + vanilla JS + sklearn + RAG over 49 PDF pages. "
                "No cloud bill. 24-hour achievable; production-plausible.",
            ],
        ],
    )

    # ── 4. Real-world ────────────────────────────────────────────────────────
    add_heading(doc, "4. Real-world impact scenarios", level=1)
    add_para(doc, "Scenario 1 — AIIMS-style ransomware at a hospital vendor", bold=True, size=12)
    add_para(
        doc,
        "A hospital's billing vendor is hit by ransomware. Under DPDP §8(5), "
        "the hospital (Data Fiduciary) is on the hook even though they never "
        "touched the malicious code. VendorGuard flags the vendor's exposed "
        "RDP + unpatched Apache + missing encryption-at-rest clause in the "
        "DPA at onboarding, scores them 'block', and blocks onboarding. "
        "Savings: ₹250 Cr statutory penalty + reputational recovery + "
        "DPO criminal liability.",
    )
    add_para(doc, "Scenario 2 — PayTrust-style payments vendor leak", bold=True, size=12)
    add_para(
        doc,
        "A fintech uses a payment-orchestration vendor whose token database "
        "leaks. Our scan of paytrust-partner.com → 11 findings, trust 12/100, "
        "₹600 Cr exposure across §5 / §8(5) / §8(6) / §8(8). When the attack "
        "fires, gateway contains it in 43 ms, WhatsApp alert goes to the "
        "CISO, CERT-In 6-hour Form-A is auto-drafted. The company's Data "
        "Protection Board response packet is ready before the 72-hour clock "
        "runs out.",
    )
    add_para(doc, "Scenario 3 — Cross-border SaaS vendor violating §16", bold=True, size=12)
    add_para(
        doc,
        "An HR SaaS vendor silently moves Indian employee data to a US-based "
        "sub-processor without the §16 cross-border transfer clause. Contract "
        "Intel v2 catches it on the DPA at signing, flags red with 0.92 "
        "confidence, quotes the gazette, offers a counter-sign rewrite. "
        "The company never signs the broken DPA.",
    )
    add_para(doc, "", size=4)
    add_para(doc, "Impact math", bold=True, size=12)
    add_para(
        doc,
        "SME with 100 vendors × 3% breach rate = 3 breaches/yr × ₹250 Cr "
        "ceiling = ₹750 Cr annual risk exposure. VendorGuard at ₹5 lakh/yr "
        "list price = 15,000× ROI if it catches just one.",
    )

    # ── 5. Architecture ──────────────────────────────────────────────────────
    add_heading(doc, "5. 5-layer architecture", level=1)
    add_table(
        doc,
        ["Layer", "What it does", "Technology"],
        [
            [
                "L1 — OSINT",
                "Pre-onboarding scan. Shodan · HIBP · crt.sh (live) · VT · DNS "
                "· TLS · nuclei.",
                "httpx async fan-out, dnspython, ssl stdlib, ProjectDiscovery "
                "nuclei subprocess.",
            ],
            [
                "L2 — Trust Score",
                "0-100 weighted composite with bands (safe ≥80, watch 50-79, "
                "block <50). All weights auditable in JSON.",
                "Python pure-function scorer.",
            ],
            [
                "L3 — DPDP Mapping",
                "15 clauses · ₹ penalty per clause · 49-passage RAG citation · "
                "ISO 27001 / SOC 2 / NIST CSF cross-walk.",
                "TF-IDF + cosine similarity + clause-reference boost (v3.2.1).",
            ],
            [
                "L4 — Runtime Gateway",
                "IsolationForest + deterministic rules · <100 ms containment · "
                "canary tripwires · SSE + webhook dispatch.",
                "scikit-learn IsolationForest (50 baseline traces) + rule "
                "engine + aiosqlite persistence.",
            ],
            [
                "L5 — Contract Intel",
                "16 rules · evidence trace (keyword+offset+snippet+confidence) "
                "· red/amber/green verdict · Act quote · rewrite.",
                "Pure-Python rule engine + regex evidence matcher + RAG quote "
                "pull-through.",
            ],
        ],
    )

    # ── 6. Tech stack ────────────────────────────────────────────────────────
    add_heading(doc, "6. Tech stack — languages, frameworks, services", level=1)
    add_para(doc, "Programming languages", bold=True, size=12)
    add_table(
        doc,
        ["Language", "Where", "% of codebase"],
        [
            ["Python 3.11+", "FastAPI backend, ML, RAG, PDF, scripts", "~72%"],
            ["JavaScript (ES2022, vanilla)", "Frontend dashboard, Demo Mode, SSE, Cytoscape", "~24%"],
            ["HTML5 + Tailwind CSS", "Frontend layout & styling", "~3%"],
            ["Shell / Bash", "ffmpeg muxing, Docker entrypoints", "<1%"],
        ],
    )
    add_para(doc, "Backend", bold=True, size=12)
    add_bullets(doc, [
        "FastAPI 0.110+ (30+ endpoints, async, auto-OpenAPI).",
        "Pydantic v2 — all request/response models typed.",
        "uvicorn (dev) / gunicorn + uvicorn workers (prod).",
        "aiosqlite — async SQLite persistence for scans, gateway state, alerts, canaries.",
        "scikit-learn IsolationForest — gateway anomaly detection.",
        "scikit-learn TfidfVectorizer + cosine similarity — RAG retrieval.",
        "httpx (async) — OSINT integrations.",
        "dnspython — SPF / DMARC / DNS.",
        "ssl + socket stdlib — TLS certificate inspection.",
        "ReportLab — Board PDF and CERT-In Form-A PDF.",
        "sse-starlette — live alert stream.",
        "Optional: anthropic, openai, openrouter (LLM polish). twilio (WhatsApp).",
    ])
    add_para(doc, "Frontend", bold=True, size=12)
    add_bullets(doc, [
        "Vanilla JS (~2300-line single file) — no framework.",
        "Tailwind CSS via CDN + handcrafted utility classes.",
        "Cytoscape.js — vendor relationship graph.",
        "EventSource (native SSE) — subscribes to /alerts/stream.",
        "Keyboard shortcuts: Ctrl+Shift+D = Demo Mode · / = Ask DPDP drawer.",
    ])
    add_para(doc, "Data assets", bold=True, size=12)
    add_bullets(doc, [
        "dpdp_act_excerpts.json — 49 verbatim RAG passages (Act §§2-34 + Rules 2025 + CERT-In).",
        "dpdp_clauses.json — 15 clauses + 18 finding→clause mappings.",
        "benchmark_dpas.json — 4 canned DPAs (strong / ambiguous / weak / saas-commodity).",
        "framework_crosswalk.json — ISO 27001 / SOC 2 / NIST CSF per clause.",
        "baseline_traffic.json — 50 seeded traces for IsolationForest.",
        "demo_vendors.json — 4 pre-scanned vendors; works offline with DEMO_MODE=true.",
    ])
    add_para(doc, "Stage-day demo pipeline", bold=True, size=12)
    add_bullets(doc, [
        "Sarvam.ai Bulbul v3 TTS (speaker 'ratan', Indian English) — voiceover.",
        "Playwright async headless chromium — automated UI recording.",
        "ffmpeg libx264 + aac — mux video + voiceover into out/demo-final.mp4.",
    ])
    add_para(doc, "Deploy targets", bold=True, size=12)
    add_bullets(doc, [
        "render.yaml — one-click blueprint (free tier).",
        "fly.toml, railway.json — alternative hosting.",
        "Dockerfile + docker-compose.yml — local / VPS.",
        "frontend/vercel.json — static Vercel deploy.",
    ])

    # ── 7. What we cover that others miss ────────────────────────────────────
    add_heading(doc, "7. What we cover that others miss", level=1)
    add_numbered(doc, [
        "Rupee-quantified penalty on every single clause (not a banner).",
        "Evidence trace on every Contract Intel verdict — keyword + offset + snippet + confidence.",
        "Gazette-page RAG citation on every Ask DPDP answer.",
        "Live Defense is an actual running gateway — not a dashboard of integrations.",
        "Reproducible /selftest harness — 4/4 benchmark DPAs verify the rule engine.",
        "CERT-In 6-hour Form-A auto-drafted on every gateway block.",
        "One-click Audit ZIP — scan.json + playbook.json + alerts.json + PDFs.",
        "Significant Data Fiduciary (§10) uplift as a first-class rule.",
        "Compliance Diff — set-arithmetic delta between two historical scans.",
        "Open core, self-hostable, ₹5 lakh/yr pricing target — bottom-of-market moat.",
    ])

    # ── 8. Competitive ───────────────────────────────────────────────────────
    add_heading(doc, "8. Competitive landscape", level=1)
    add_table(
        doc,
        ["Capability", "VendorGuard", "OneTrust", "Securiti", "BigID", "Tsaaro"],
        [
            ["DPDP Act §-level coverage", "✓ 15 clauses verbatim", "GDPR-first + module", "DPDP module", "DPDP module", "Consulting"],
            ["DPDP Rules 2025 R.1-R.22", "✓ indexed", "Partial", "Partial", "Partial", "Manual"],
            ["₹ penalty per clause", "✓", "—", "—", "—", "Manual"],
            ["Evidence-traced contract rules", "✓ 16 rules", "LLM opaque", "LLM opaque", "LLM opaque", "—"],
            ["Gazette-page RAG citation", "✓ 49 passages", "—", "—", "—", "PDF appendix"],
            ["Live OSINT (crt.sh no-key)", "✓", "—", "—", "—", "—"],
            ["Sub-second behavioural gateway", "✓ 43ms", "—", "—", "—", "—"],
            ["Canary tokens", "✓ <200ms mint", "Via partner", "Via partner", "Via partner", "—"],
            ["CERT-In 6h Form-A PDF", "✓", "—", "—", "—", "Manual"],
            ["Self-test harness", "✓ 4/4", "—", "—", "—", "—"],
            ["Self-hostable / open core", "✓ MIT", "—", "—", "—", "—"],
            ["SME price (200 vendors)", "₹2-5 L/yr", "₹40-65 L", "₹30-50 L", "₹50-65 L", "₹12-30 L one-time"],
        ],
    )
    add_para(
        doc,
        "We are not claiming feature-parity with a $80K OneTrust SKU. We are "
        "claiming fit-for-purpose at the India SME-MSME tier they don't reach.",
        italic=True,
        color=MUTED,
    )

    # ── 9. Business model ────────────────────────────────────────────────────
    add_heading(doc, "9. Business model — how we make money", level=1)
    add_para(doc, "Pricing ladder", bold=True, size=12)
    add_table(
        doc,
        ["Tier", "Audience", "Vendors", "₹ / year", "Differentiator"],
        [
            ["Community (open core)", "Solo / dev / students", "Unlimited, self-host", "₹0", "Full product, MIT"],
            ["Starter (cloud, shared)", "Startups (1-50)", "50", "₹1.5 L", "Managed hosting, email support"],
            ["Growth (cloud, single tenant)", "SMEs (50-250)", "250", "₹4.8 L", "SSO, RBAC, WhatsApp, CERT-In auto"],
            ["Enterprise (self-host / VPC)", "BFSI / healthtech / govt", "Unlimited", "₹25 L+", "On-prem, custom LLM, SLAs"],
            ["Consulting add-on", "All tiers", "—", "₹50K / engagement", "DPO-as-a-service, audit prep"],
        ],
    )
    add_para(doc, "Revenue mix (Year 2)", bold=True, size=12)
    add_bullets(doc, [
        "70% subscription ARR (Starter + Growth + Enterprise).",
        "20% consulting (audit prep, DPO-as-a-service, DPA templating).",
        "10% training + certification (Certified VendorGuard Analyst programme).",
    ])
    add_para(doc, "Unit economics (Growth tier)", bold=True, size=12)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["ACV", "₹4.8 L"],
            ["Gross margin", "~92%"],
            ["CAC (LinkedIn + DPO community)", "₹40K"],
            ["CAC payback", "~1 month"],
            ["Net-revenue retention target", "115%"],
        ],
    )
    add_para(doc, "Market sizing", bold=True, size=12)
    add_bullets(doc, [
        "~2 million digital MSMEs × 1% regulated = 20,000 SAM.",
        "× ₹3 L average ACV = ₹600 Cr ARR (SME segment).",
        "Enterprise: ~2,000 logos × ₹25 L = ₹500 Cr ARR.",
        "Conservative Year-3 target: 100 logos × ₹3 L = ₹3 Cr ARR.",
    ])

    # ── 10. Killer lines ─────────────────────────────────────────────────────
    pagebreak(doc)
    add_heading(doc, "10. Killer pitch lines (memorize these)", level=1)
    for label, line in [
        (
            "Opening hook (8s)",
            "Every enterprise in India runs on vendors, and under the DPDP "
            "Act any one vendor breach is YOUR 250-crore problem. "
            "VendorGuard catches it in 43 milliseconds.",
        ),
        (
            "The differentiator (6s)",
            "OneTrust tells you a clause is risky. We show you the keyword, "
            "the offset, the gazette page, and the rewrite — evidence, not "
            "vibes.",
        ),
        (
            "The moat (8s)",
            "DPDP-native, not GDPR-retrofitted. Rules where rules belong. "
            "ML where anomalies live. LLM only for polish. Nothing hidden.",
        ),
        (
            "The urgency (5s)",
            "DPDP Rules 2025 are landing now. The first mover who ships "
            "evidence-grounded audit packs wins the next 18 months.",
        ),
        (
            "The close (5s)",
            "Scan, score, map, defend, diff, audit. One screen, one ZIP, "
            "one rupee number that keeps the CISO employed. That's "
            "VendorGuard.",
        ),
    ]:
        add_para(doc, label, bold=True, size=12, color=BRAND)
        add_para(doc, f"\u201c{line}\u201d", italic=True, size=12)
        add_para(doc, "", size=4)

    add_para(
        doc,
        "Total hook + close = 32 seconds, leaves 58 seconds for live demo "
        "within a 90-second slot.",
        italic=True,
        color=MUTED,
    )

    # ── 11. 90-second delivery ───────────────────────────────────────────────
    add_heading(doc, "11. 90-second pitch delivery script", level=1)
    add_para(
        doc,
        "Each step is 10-12 seconds of speaking. Click in the UI as you speak "
        "— the screen reinforces every sentence.",
        italic=True,
        color=MUTED,
    )
    steps = [
        ("Executive Board",
         "'One screen for the CISO. Four vendors tracked, average trust 34, "
         "₹1,550 Cr combined exposure, one attack blocked this week, ₹600 Cr "
         "saved. ISO 27001 / SOC 2 / NIST CSF coverage mapped automatically.'"),
        ("Vendor Scan → Findings",
         "'Click paytrust. Real OSINT — Shodan, HIBP, crt.sh, VirusTotal, DNS, "
         "TLS, nuclei. Every finding is timestamped, sourced, and carries a "
         "verify URL you can open yourself.'"),
        ("Vendor Scan → DPDP",
         "'Every finding maps to a verbatim DPDP clause — §5, §8(5), §8(6), "
         "§8(8) — with rupee penalty and ISO / SOC 2 / NIST cross-walk.'"),
        ("Contract Intel → Benchmark",
         "'Click the Weak DPA chip. 16 deterministic rules fire in under a "
         "second: 9 red, 4 amber, ₹1,650 Cr exposure. Each verdict has a "
         "confidence score, evidence trace with keyword + snippet, Act "
         "quote, ready-to-counter-sign rewrite.'"),
        ("Live Defense → Gateway",
         "'IsolationForest + rules. Simulate an exfiltration. Sub-100 "
         "millisecond containment. Ticker bumps blocks + ₹ saved.'"),
        ("Remediation → Playbook / Diff / Audit ZIP",
         "'Monday checklist — owner, SLA, ₹ savings. Diff sub-tab shows what "
         "changed since last scan. Audit ZIP button packages scan + playbook "
         "+ alerts + board PDF + CERT-In Form-A. Hand to the Data Protection "
         "Board in one click.'"),
        ("Ask DPDP",
         "'Press slash. RAG over the 49-passage corpus. Ask \"penalties "
         "under 8(5)\" — the answer quotes the gazette directly with a page "
         "citation. No hallucination.'"),
        ("Rule-engine selftest",
         "'Click the selftest chip. All 4 benchmark DPAs verify in real "
         "time. 4 out of 4. Reproducible. Not a black box.'"),
    ]
    for i, (label, line) in enumerate(steps, start=1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.color.rgb = BRAND
        p.add_run(line)

    add_para(
        doc,
        "Closer: 'Rules where rules belong. ML where anomalies live. LLM "
        "only for polish. Nothing hidden. That's why VendorGuard wins.'",
        bold=True,
        size=12,
    )

    # ── 12. Judge Q&A ────────────────────────────────────────────────────────
    pagebreak(doc)
    add_heading(doc, "12. Judge Q&A — 15 hard questions with answers", level=1)
    qa = [
        ("Where is the actual AI? Sounds like regex + dashboards.",
         "Three clearly labelled layers. (1) IsolationForest on 50 behavioural "
         "baselines for the gateway — real unsupervised anomaly detection. "
         "(2) TF-IDF + cosine similarity over 49 DPDP passages for Ask DPDP — "
         "verifiable retrieval, not generative. (3) Optional LLM polish for "
         "the Executive summary, grounded in RAG and labelled as such in the "
         "UI. We chose rules for legal clauses on purpose — a hallucinated "
         "DPA finding is a lawsuit."),
        ("Isn't this just OneTrust / Securiti?",
         "No. DPDP-native taxonomy (not GDPR-retrofitted). Evidence-trace "
         "auditable (keyword + offset + snippet + confidence on every "
         "verdict). Sub-second runtime gateway with CERT-In 6-hour Form-A "
         "auto-drafted — a category OneTrust doesn't play in. And priced for "
         "the India SME segment they structurally can't reach."),
        ("How do we know your 16 contract rules are correct?",
         "Click the selftest chip. Runs all 4 benchmark DPAs through the "
         "rule engine and verifies every verdict matches the expected one "
         "baked into benchmark_dpas.json. curl /selftest returns the same "
         "JSON. Reproducible. Auditable."),
        ("Your 43ms containment — isn't that theatre?",
         "time.perf_counter() around the full /gateway/proxy handler. "
         "Published in the alert payload as containment_seconds. alert_id "
         "committed to SQLite before the response returns — verifiable."),
        ("What if the LLM / Shodan / HIBP APIs go down on stage?",
         "DEMO_MODE=true pre-seeds 4 demo vendors, no keys required. "
         "crt.sh / DNS / TLS are live but keyless. Stage-day MP4 is the "
         "final backstop — if the projector plays MP4 we can still demo."),
        ("Who is the paying customer — CISO, DPO, procurement, or legal?",
         "All four, one screen. CISO: trust ring + containment + ₹ saved. "
         "DPO: clause map + gazette citations + CERT-In PDF. Procurement: "
         "DPA red/amber/green + rewrite + confidence. Legal: Audit ZIP. "
         "One product replaces a three-tool stack."),
        ("What if DPDP Rules 2025 get amended next year?",
         "RAG corpus, rule catalog, clause map are all JSON/Python data "
         "files. Drop in new passages, add a rule, ship. Catalog is "
         "versioned so Audit ZIP records the exact catalog hash the verdict "
         "was signed against."),
        ("Why not train a custom LLM on DPDP?",
         "Legal citation demands verbatim section + gazette page, not a "
         "model paraphrase. TF-IDF + verbatim quotes is more defensible "
         "than a fine-tuned LLM. Scaling to BM25 or a local embedding model "
         "(bge-small-en-v1.5) is a 2-day upgrade if we ever need it."),
        ("How does this compare to Tsaaro / Lexplosion / Cerebrus?",
         "They are consulting shops. They ship PDFs, templates, and partner "
         "manhours. We ship a runtime product — scan, score, gateway, "
         "audit ZIP. Partnership is a realistic go-to-market wedge (they "
         "use us as a DPA scoring tool), not head-to-head competition."),
        ("Privacy of the scan data itself?",
         "Self-hostable. A BFSI enterprise runs the entire stack inside "
         "their VPC — single FastAPI process + SQLite file. Structural "
         "advantage over SaaS incumbents that must route through US/EU "
         "clouds."),
        ("Why a gateway on top of a contract scanner? Scope creep?",
         "Contract = the promise, gateway = the enforcement. DPDP §8(5) "
         "requires the Data Fiduciary to actually ensure safeguards, not "
         "just document them. ADT argument: contract = sign on the lawn, "
         "gateway = the alarm that actually fires."),
        ("What does a breach actually cost under DPDP?",
         "Up to ₹250 crore per type of contravention (Schedule). One "
         "breach typically triggers 3-5 contraventions simultaneously. "
         "Combined exposure per incident: ₹500 Cr – ₹1,250 Cr. Plus "
         "reputation + DPO criminal liability. VendorGuard at ₹5 lakh/yr "
         "is a 1,000× risk-reward ratio even before litigation."),
        ("What's your team background?",
         "Undergraduates from DSCE and BMSCE. Built this in ~30 hours "
         "from a master prompt. The code in this repo is what we shipped "
         "— 32 pytest cases, 30+ endpoints, 5-panel dashboard, "
         "Playwright-recorded demo video. Happy to walk through any "
         "module live."),
        ("Business model?",
         "Open core. MIT community edition (full functionality, self-host). "
         "Paid tiers — Starter ₹1.5 L, Growth ₹4.8 L, Enterprise ₹25 L+ "
         "— for managed hosting + SSO / RBAC / audit logs / on-prem "
         "support. Consulting revenue on top. Year-3 conservative target: "
         "₹3 Cr ARR."),
        ("What happens in the next 6 months if you win?",
         "(1) Deploy backend to Render. (2) Onboard 5 Bengaluru DPO design "
         "partners. (3) Ship clause-pack subscription (BFSI, healthtech, "
         "edtech). (4) Law-firm co-brand — Cyril Amarchand / Khaitan / "
         "Nishith Desai logo on the Audit ZIP. (5) Grow RAG corpus to 200 "
         "passages (Rules 2025 final text, sectoral guidance, RBI/SEBI "
         "circulars)."),
    ]
    for i, (q, a) in enumerate(qa, start=1):
        add_para(doc, f"Q{i}. {q}", bold=True, size=11.5, color=BRAND)
        add_para(doc, f"A. {a}", size=11)
        add_para(doc, "", size=4)

    # ── 13. Glossary ─────────────────────────────────────────────────────────
    pagebreak(doc)
    add_heading(doc, "13. Glossary — every abbreviation + full form", level=1)
    glossary = [
        ("DPDP", "Digital Personal Data Protection — the 2023 Indian Act + 2025 Rules we map to."),
        ("DPB", "Data Protection Board of India — regulator created under DPDP."),
        ("DPO", "Data Protection Officer — India-mandated role for Significant Data Fiduciaries."),
        ("SDF", "Significant Data Fiduciary — DPDP §10 designation with uplifted obligations."),
        ("DPA", "Data Processing Agreement — contract between Data Fiduciary and Data Processor."),
        ("PII", "Personally Identifiable Information — the data type DPDP protects."),
        ("GDPR", "General Data Protection Regulation — EU; DPDP is often compared to it."),
        ("CCPA", "California Consumer Privacy Act — US state-level privacy law."),
        ("CERT-In", "Indian Computer Emergency Response Team — requires 6-hour incident reporting."),
        ("RBAC", "Role-Based Access Control — Enterprise tier."),
        ("SSO", "Single Sign-On — Enterprise tier."),
        ("SIEM", "Security Information and Event Management — category we complement (Splunk, Elastic)."),
        ("OSINT", "Open Source Intelligence — what L1 is (Shodan, HIBP, crt.sh, VT, DNS, TLS)."),
        ("HIBP", "HaveIBeenPwned — breach corpus lookup."),
        ("VT", "VirusTotal — domain/URL reputation."),
        ("CT / crt.sh", "Certificate Transparency logs — live subdomain enumeration (no key)."),
        ("TLS", "Transport Layer Security — HTTPS; we probe it."),
        ("SPF / DMARC", "Sender Policy Framework / DMARC — DNS email-auth records."),
        ("RAG", "Retrieval-Augmented Generation — our DPDP answer pattern (retrieve + quote)."),
        ("TF-IDF", "Term Frequency × Inverse Document Frequency — retrieval algorithm."),
        ("ML", "Machine Learning — IsolationForest at the gateway."),
        ("LLM", "Large Language Model — Claude / GPT, only for Executive summary polish."),
        ("SSE", "Server-Sent Events — live alert stream."),
        ("TTS", "Text-to-Speech — Sarvam.ai for demo voiceover."),
        ("API", "Application Programming Interface — the FastAPI layer."),
        ("SaaS", "Software as a Service — incumbent distribution model."),
        ("VPC", "Virtual Private Cloud — Enterprise self-host option."),
        ("ACV", "Annual Contract Value."),
        ("ARR", "Annual Recurring Revenue."),
        ("CAC", "Customer Acquisition Cost."),
        ("SLA", "Service Level Agreement — Playbook SLA (7 / 30 / long days)."),
        ("GRC", "Governance, Risk & Compliance — the OneTrust / Securiti category."),
        ("KPI", "Key Performance Indicator — Executive Board metrics."),
        ("CVE", "Common Vulnerabilities and Exposures — nuclei CVE templates."),
        ("MVP", "Minimum Viable Product — what we shipped in 30 hours."),
        ("SME / MSME", "Small & Medium Enterprise / Micro, Small & Medium Enterprise — target segment."),
        ("BFSI", "Banking, Financial Services, Insurance — vertical with highest DPDP exposure."),
        ("CDP", "Chrome DevTools Protocol — how Playwright drives our demo recording."),
    ]
    add_table(doc, ["Short", "Full form + context"], [[s, f] for s, f in glossary])

    # ── 14. Team knowledge ───────────────────────────────────────────────────
    pagebreak(doc)
    add_heading(doc, "14. Team knowledge base & stage-day checklist", level=1)
    add_para(doc, "Critical numbers — every speaker must know these cold", bold=True, size=12)
    add_table(
        doc,
        ["Fact", "Number / Value"],
        [
            ["DPDP Act enacted", "11 Aug 2023"],
            ["DPDP Rules 2025", "R.1 – R.22"],
            ["Max penalty per contravention", "₹250 crore"],
            ["Breach SLA to DPB", "72 hours (§8(6))"],
            ["CERT-In incident SLA", "6 hours (2022 directions)"],
            ["RAG corpus passages", "49"],
            ["Contract Intel rules", "16"],
            ["DPDP clauses mapped", "15"],
            ["Benchmark DPAs", "4 (strong / ambiguous / weak / saas-commodity)"],
            ["Gateway containment time (paytrust)", "43 ms wall-clock"],
            ["Trust score bands", "safe ≥80 · watch 50-79 · block <50"],
            ["pytest", "31 / 32 passing (1 pre-existing flake)"],
            ["Self-test", "4 / 4 ✓"],
            ["Demo vendors", "4 pre-scanned, no keys required"],
            ["Stage-day MP4 length", "2 min 40 sec"],
        ],
    )
    add_para(doc, "Ownership matrix", bold=True, size=12)
    add_table(
        doc,
        ["Domain", "Primary", "Backup"],
        [
            ["90-second pitch delivery", "Captain", "Backup speaker"],
            ["Live demo driving (keyboard)", "Captain", "Dev"],
            ["Q&A on ML / gateway", "Dev", "Captain"],
            ["Q&A on DPDP / legal", "DPDP lead", "Captain"],
            ["Q&A on business model", "Captain", "Any"],
            ["Stage-day MP4 backup", "Captain's laptop", "USB stick"],
            ["Deploy backend to Render", "Dev", "—"],
            ["Rehearse pitch 10×", "All", "—"],
        ],
    )
    add_para(doc, "Stage-day checklist", bold=True, size=12)
    add_bullets(doc, [
        "Backend deployed to Render — hit /health once to warm the cold start.",
        "Frontend deployed to Vercel; ?api= points at the live backend.",
        "out/demo-final.mp4 on Captain's laptop + USB stick.",
        "Ctrl+Shift+D tested on venue projector resolution.",
        "Judge Q&A printout (section 12) in hand.",
        "Pitch rehearsed 10 times, timed under 90 seconds.",
        "Opening + closing lines memorized verbatim.",
        "selftest chip shows 4/4 ✓ in the header before pitch starts.",
    ])
    add_para(doc, "Failure modes & recovery", bold=True, size=12)
    add_table(
        doc,
        ["Failure", "Recovery"],
        [
            ["Stage wifi dies", "Play out/demo-final.mp4; Captain narrates over the video."],
            ["Render cold start >10s", "Frontend falls back to localStorage-cached demo vendor data; still a full demo."],
            ["Ask DPDP returns wrong section", "Move on; mention PR #3 boost covers §8(5)."],
            ["Judge asks about production scale", "'Hackathon MVP. Next 6 months: Render prod + 5 design partners.'"],
            ["Live backend not deployed in time", "Run locally; frontend points at local IP; enable MP4 backup."],
            ["Gateway test fails on stage", "'Known ML baseline flake across /gateway/reset — passes in isolation. Focus on the 31 green.'"],
        ],
    )

    # ── 15. Changelog ────────────────────────────────────────────────────────
    pagebreak(doc)
    add_heading(doc, "15. Changelog (v2 → v3.2.1)", level=1)
    add_para(doc, "v3.2.1 — cold-read polish + stage-day demo assets", bold=True, size=12)
    add_bullets(doc, [
        "RAG clause-reference boost: 'penalties under §8(5)' → returns §8(5), not §32.",
        "14 → 49 copy drift fix in Ask DPDP subtitle, Board Report footer, README.",
        "Header version chip aligned with __version__ and /health — all read v3.2.1.",
        "Stage-day demo assets: docs/DEMO_SCRIPT.md + scripts/generate_voiceover.py "
        "+ scripts/record_demo.py + out/demo-voiceover.mp3 + out/demo-final.mp4.",
        "New pytest case test_rag_ask_boosts_explicit_section_reference.",
    ])
    add_para(doc, "v3.2 — Audit ZIP + Compliance Diff + CSV + selftest", bold=True, size=12)
    add_bullets(doc, [
        "One-click DPDP Audit Evidence Bundle (/audit/{vendor}.zip).",
        "Compliance Diff sub-tab + /scan/{v}/diff.",
        "Playbook CSV export with Jira/Linear/GitHub Projects headers.",
        "Fourth benchmark DPA (saas-commodity-dpa).",
        "Demo mode v2 (Weak DPA chip + Compliance Diff sub-tab).",
        "Rule-engine /selftest endpoint + header chip.",
        "27 → 31 pytest cases.",
    ])
    add_para(doc, "v3.1 — sidebar consolidation + rule engine v2", bold=True, size=12)
    add_bullets(doc, [
        "13-item sidebar → 5 top-level panels.",
        "Contract Intel v2: 12 → 16 rules; confidence + evidence_trace.",
        "Benchmark Evidence Ledger (3 canned DPAs; grew to 4 in v3.2).",
        "Live crt.sh / DNS / TLS always-on; /osint/live/{vendor}.",
        "DPDP RAG corpus 14 → 49 passages.",
        "Compliance Diff backend (scan_history + /scan/{v}/diff).",
        "PITCH.md + COMPETITIVE.md.",
    ])
    add_para(doc, "v3.0 — Layer 5 + Executive Board + Playbook", bold=True, size=12)
    add_bullets(doc, [
        "Contract Intelligence /contract/analyze with 12 rules.",
        "Executive Board /portfolio + /kpis.",
        "Monday Playbook grouped by DPDP clause.",
        "Framework cross-walk (ISO/SOC2/NIST CSF).",
        "CERT-In 6h Form-A PDF /incident/{alert_id}.pdf.",
        "Bulk onboarding /vendors/bulk.",
        "Generic webhook alerts.",
        "Floating Ask DPDP drawer (/ hotkey).",
        "Demo Mode (Ctrl+Shift+D).",
    ])
    add_para(doc, "v2.0 — foundation", bold=True, size=12)
    add_bullets(doc, [
        "IsolationForest gateway + anomaly_score on alerts.",
        "Real containment time (time.perf_counter).",
        "ProjectDiscovery nuclei subprocess runner.",
        "SQLite persistence via aiosqlite.",
        "Board PDF report (ReportLab).",
        "DPDP Act RAG (TF-IDF).",
        "SSE dashboard + vendor graph + canary tokens.",
        "4 demo vendors, 14-case pytest, Docker / Fly / Railway / Vercel ready.",
    ])

    # ── 16. Credits ──────────────────────────────────────────────────────────
    add_heading(doc, "16. Credits & license", level=1)
    add_para(
        doc,
        "Built for Athernex 2026 (DSCE × BMSCE) by Team Rashi Innovators. "
        "Open-source credits: FastAPI (Sebastián Ramírez), scikit-learn, "
        "httpx (Encode), ReportLab, Playwright (Microsoft), Cytoscape.js, "
        "Tailwind CSS, ProjectDiscovery nuclei, crt.sh (Sectigo), Sarvam.ai "
        "(TTS). MIT licensed — use freely, fork freely, ship anything.",
    )
    add_para(doc, "", size=6)
    add_para(
        doc,
        "Rules where rules belong. ML where anomalies live. LLM only for "
        "polish. Nothing hidden. That's why VendorGuard wins.",
        bold=True,
        italic=True,
        size=13,
        color=BRAND,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
